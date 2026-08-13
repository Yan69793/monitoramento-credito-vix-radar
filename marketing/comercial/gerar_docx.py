#!/usr/bin/env python3
"""
Gera documento Word com resumo detalhado do Plano Comercial VIX Radar Fase 1.
Formato: bullets points, indices, tabelas, secoes numeradas.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Style setup ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.space_before = Pt(0)

for level in range(1, 4):
    heading = doc.styles[f'Heading {level}']
    heading.font.name = 'Calibri'
    heading.font.color.rgb = RGBColor(0x00, 0x10, 0x20)
    if level == 1:
        heading.font.size = Pt(18)
        heading.paragraph_format.space_before = Pt(24)
        heading.paragraph_format.space_after = Pt(12)
    elif level == 2:
        heading.font.size = Pt(14)
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(8)
    else:
        heading.font.size = Pt(12)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

def add_bullet(text, bold_prefix=None, level=0):
    """Add a bullet point, optionally with bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(headers, rows):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()  # spacer
    return table

def add_tag(text, color_hex="001020"):
    """Add a tag/badge style line (like [Fato], [Risco], etc.)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(
        int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
    )
    return p

# ══════════════════════════════════════════════════════════════
# CAPA
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('VIX Radar')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x00, 0x10, 0x20)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Plano Comercial — Fase 1')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0xB7, 0x98, 0x5D)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Documento interno. Szuchmacher Consultoria Ltda.\n30 de julho de 2026').font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# SUMARIO / INDICE
# ══════════════════════════════════════════════════════════════
doc.add_heading('Sumario', level=1)
toc_items = [
    '0. Decisao principal — contratar agencia agora ou vender sozinho primeiro',
    '1. Bloqueios P0 antes de qualquer verba de midia',
    '2. Revisao critica do posicionamento existente',
    '3. Ofertas comercializaveis dentro da estrutura atual',
    '4. Meta numerica do piloto de 90 dias',
    '5. Pesquisa de agencias qualificadas',
    '6. Ranking e o problema do orcamento',
    '7. Roteiro de abordagem para agencias',
    '8. Briefing padronizado (arquivo separado)',
    '9. Grade de comparacao de propostas',
    '10. Orcamento estimado total',
    '11. Riscos',
    '12. O que depende de voce para destravar a Fase 2',
    '13. Testes de sensibilidade',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 0. DECISAO PRINCIPAL
# ══════════════════════════════════════════════════════════════
doc.add_heading('0. Decisao principal', level=1)

p = doc.add_paragraph()
run = p.add_run('Contratar agencia agora.')
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph(
    'Decisao do operador. A rota solo foi considerada e descartada. O caminho escolhido e '
    'contratacao de agencia para execucao de ABM, outbound e midia sobre a estrategia que ja existe.'
)

doc.add_heading('Tres condicoes para o dinheiro nao queimar', level=2)
add_bullet('Cobranca funcionando antes do primeiro anuncio. Sem checkout, a agencia entrega reuniao que nao vira receita.', bold_prefix='1. ')
add_bullet('Orcamento calibrado ao que o mercado aceita. Agencia full-service (Sabiá, Intelligenzia) nao opera na faixa de R$ 4.000/mes. O caminho realista e boutique de outbound, consultor de ABM independente ou SDR terceirizado.', bold_prefix='2. ')
add_bullet('Piloto de 90 dias com clausula de saida. Sem contrato longo antes de validar execucao.', bold_prefix='3. ')

doc.add_heading('Ordem de execucao', level=2)
add_bullet('Resolver P0-1 (cobranca) e P0-3 (confirmar pagantes) antes de assinar com agencia.')
add_bullet('O resto dos P0 pode ser paralelo a busca e contratacao.')
add_bullet('Se o orcamento real for menor que o estimado, descer para consultor independente em vez de agencia.')

p = doc.add_paragraph()
run = p.add_run('Risco principal: ')
run.bold = True
p.add_run('queimar verba de midia e fee de agencia para gerar reunioes que nao convertem '
           'porque nao existe checkout nem processo de cobranca.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. BLOQUEIOS P0
# ══════════════════════════════════════════════════════════════
doc.add_heading('1. Bloqueios P0 antes de qualquer verba de midia', level=1)

doc.add_paragraph(
    'Nove bloqueios listados em ordem de prioridade. Nenhuma agencia resolve os tres primeiros. '
    'Os itens 5, 6 e 7 ja constavam como pendencia desde 13/07 e continuam nao cumpridos.'
)

add_table(
    ['#', 'Bloqueio', 'Esforco', 'Quem resolve', 'Impacto se nao resolver'],
    [
        ['P0-1', 'Cobranca automatizada (Stripe/MercadoPago)', '3-5 dias dev', 'Desenvolvedor',
         'Cliente diz sim e nao consegue pagar'],
        ['P0-2', 'Processo manual de contrato, emissao e ativacao', '1 dia', 'Yan',
         'Ponte ate o P0-1 ficar pronto'],
        ['P0-3', 'Confirmar numero real de clientes pagantes', '30 min', 'Yan',
         'Dimensiona ou cancela todo o resto'],
        ['P0-4', 'Remover case Mirabaud do material comercial', '2h', 'Yan + dev',
         'Contraparte judicial citada como cliente'],
        ['P0-5', 'LinkedIn Insight Tag em vixradar.com', '30 min', 'Desenvolvedor',
         'Sem retargeting, sem Matched Audiences'],
        ['P0-6', 'Pagina institucional VIX Radar no LinkedIn', '1h', 'Yan',
         'Sem lastro para anuncios e Document Ads'],
        ['P0-7', 'Landing dedicada de captura da newsletter', '2h', 'Desenvolvedor',
         'Topo de funil sem ponto de conversao'],
        ['P0-8', 'Reprocessar noturno 30/07', 'Executar script', 'Yan',
         'Dados do dia sob deepseek-v4-flash, nao Anthropic'],
        ['P0-9', 'CRM definido e operacional', '2-4h', 'Yan',
         'Sem tracking de funil'],
    ]
)

add_tag('[Fato] Worker v4.9.183 com ok:true, verificador_ok:true em 30/07 21:55 UTC (18:55 BRT).')
add_tag('[Fato] Pre-requisitos nao cumpridos desde 13/07: Insight Tag, pagina no LinkedIn, landing de newsletter, Sales Navigator.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. REVISAO DE POSICIONAMENTO
# ══════════════════════════════════════════════════════════════
doc.add_heading('2. Revisao critica do posicionamento', level=1)

doc.add_heading('Solido — nao mexer', level=2)
add_bullet('ICP com anti-ICP: universo real de 8-15 mil profissionais, segmentacao por Skill.', bold_prefix='ICP cirurgico. ')
add_bullet('"Saber do evento antes do mercado precificar", "parar de varrer CVM na mao". Frases que um gestor reconhece como dele.', bold_prefix='Jobs-to-be-done validados. ')
add_bullet('Ordem correta: reconhecer forca do terminal caro, explicar razao estrutural do escopo mais estreito, preco como consequencia.', bold_prefix='Contraposicionamento de preco. ')
add_bullet('Gancho Story/Statement, numeros exatos, CTA por estagio de funil. Calibrado com AuthoredUp (309k posts).', bold_prefix='Regras de copy. ')
add_bullet('Cronograma, split de orcamento, CPL alvo calibrado ao ticket real.', bold_prefix='Plano de campanha de 8 semanas. ')
add_bullet('Concorrentes, frameworks, dados de engajamento, jargao de credito. Fundamentacao acima da media.', bold_prefix='Pesquisa de mercado. ')
add_bullet('10 posts e 3 anuncios prontos e revisados.', bold_prefix='Conteudo pronto. ')
add_bullet('Dark luxury, navy + gold, paleta definida.', bold_prefix='Briefing visual. ')

doc.add_heading('Precisa de ajuste — cinco pontos', level=2)

add_bullet('O diretorio Processo Mirabaud contem acao trabalhista "Yan Szuchmacher x Mirabaud Investimentos Ltda." com dossie, calculo de verbas e provas de subordinacao. E contraparte judicial, nao cliente. O nome Mirabaud precisa sair de product.md, da apresentacao institucional, do deck Bradesco e de toda copy futura. ', bold_prefix='Case ancora Mirabaud — REMOVER. ')
add_bullet('O deck cita 17 usuarios sem separar pagantes de demo/cortesia. Com zero pagantes, o numero e enganoso. Trocar por descricao qualitativa ou remover.', bold_prefix='Track record de 17 usuarios. ')
add_bullet('O incidente de roteamento de 30/07 provou que a alegacao "IA Anthropic Claude" pode ser falsa sem deteccao. Adicionar defesa interna de proveniencia. Nao publicar, e defesa para due diligence.', bold_prefix='Alegacao de IA. ')
add_bullet('Operador rejeitou duas versoes de copy por soarem "comerciais demais". Preferir voz de analista/observacao tecnica, evitar primeira pessoa de fundador com cadencia de pitch.', bold_prefix='Voz do Yan vs voz comercial. ')
add_bullet('Estrategia certa, mas landing dedicada nao existe. Sem ela, trafego pago para newsletter desperdica verba.', bold_prefix='Newsletter como topo de funil. ')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. OFERTAS COMERCIALIZAVEIS
# ══════════════════════════════════════════════════════════════
doc.add_heading('3. Ofertas comercializaveis', level=1)

doc.add_heading('Essencial — R$ 119/mes', level=2)
add_bullet('103 emissores, 13 setores, eventos ranqueados por materialidade, EWS, watchlist, agenda CVM, relatorio PDF.')
add_bullet('Publico: analistas de credito, gestores de asset media, family offices, RI de emissores.')
add_bullet('Produto pronto. Cobranca nao. O botao "Assinar" abre formulario de cadastro, nao checkout.', bold_prefix='Status: ')

doc.add_heading('Profissional — R$ 490/mes', level=2)
add_bullet('Tudo do Essencial + briefing executivo, historico estendido, multiplos usuarios.')
add_bullet('Publico: gestoras com mais de um analista, tesourarias corporativas, comites de investimento.')
add_bullet('Mesmo bloqueio de cobranca do Essencial.', bold_prefix='Status: ')

doc.add_heading('Newsletter gratuita — topo de funil', level=2)
add_bullet('Boletim semanal com eventos mais materiais, sem paywall. Double opt-in via Resend, unsubscribe HMAC.')
add_bullet('Sistema de envio existe. Landing dedicada de captura nao.', bold_prefix='Status: ')

doc.add_heading('Tier institucional / RPPS — a construir', level=2)
add_bullet('Segmento comprovadamente pagante: concorrentes cobram R$ 1.940 a 2.810/licenca/mes em contrato publico verificado.')
add_bullet('Maior oportunidade de receita por conta. Preco sugerido: R$ 1.500 a 2.500/mes por licenca.')
add_bullet('Exige preparacao documental de 30-60 dias: certidoes, atestado de capacidade tecnica, contrato social, proposta em formato de licitacao.')
add_bullet('Custo: zero financeiro. So tempo do Yan.')
add_bullet('Verificar se o CNPJ atual comporta venda para governo.', bold_prefix='[Validar] ')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. META DO PILOTO
# ══════════════════════════════════════════════════════════════
doc.add_heading('4. Meta numerica do piloto de 90 dias', level=1)

doc.add_heading('Premissas (nao validadas)', level=2)
add_table(
    ['Variavel', 'Valor', 'Fonte'],
    [
        ['Clientes pagantes atuais', '0 (premissa)', '[Validar]'],
        ['Ticket medio mensal', 'R$ 250 (mix Essencial/Profissional)', 'Landing'],
        ['LTV estimado (12 meses)', 'R$ 1.800', 'Estimativa com churn'],
        ['Capacidade de reunioes/semana', '4', '[Validar]'],
        ['Orcamento total 90 dias', 'R$ 15.000 a 30.000', '[Validar]'],
        ['CAC toleravel', 'Ate R$ 600 (3 meses de receita)', 'Derivado do ticket'],
    ]
)

doc.add_heading('Cenario com agencia — rota escolhida', level=2)
add_bullet('Agencia ou consultor executa ABM, outbound e midia. Yan atende as reunioes geradas. Custo total: R$ 21.390.')
add_table(
    ['Mes', 'Acao', 'Reunioes', 'Oportunidades', 'Contratos'],
    [
        ['1', 'Setup CRM, listas ABM, inicio outbound, anuncios', '12', '5', '2'],
        ['2', 'Otimizacao, nutricao, follow-up', '14', '7', '3'],
        ['3', 'Escala do que funcionou, referral', '16', '8', '4'],
    ]
)
p = doc.add_paragraph()
run = p.add_run('Meta 90 dias: 9 clientes pagantes, receita ~R$ 2.250/mes, CAC R$ 2.377.')
run.bold = True
doc.add_paragraph(
    'Este CAC e viavel para o Profissional (recupera em 5 meses) mas inviavel para o Essencial '
    '(20 meses). A operacao precisa mirar o ticket de R$ 490 ou o tier RPPS.'
)

doc.add_heading('Cenario solo — referencia (descartado pelo operador)', level=2)
add_bullet('Yan opera o funil sobre a rede pessoal. Custo total em 90 dias: R$ 390 (apenas Sales Navigator).')
add_table(
    ['Mes', 'Acao', 'Reunioes', 'Oportunidades', 'Contratos'],
    [
        ['1', 'Networking + LinkedIn organico + follow-up', '8', '3', '1'],
        ['2', 'Continua + newsletter para contatos frios', '8', '4', '2'],
        ['3', 'Continua + referral dos primeiros clientes', '8', '5', '2'],
    ]
)
p = doc.add_paragraph()
run = p.add_run('Meta 90 dias: 5 clientes pagantes, receita recorrente ~R$ 1.250/mes, CAC R$ 78.')
run.bold = True
doc.add_paragraph(
    'Este CAC e inviavel para o Essencial (20 meses para recuperar). So se justifica '
    'com mix pesado em Profissional (R$ 490) e contratos de 12+ meses.'
)

doc.add_heading('Limite de CPL por ticket', level=2)
add_table(
    ['Ticket', 'CPL maximo para nao operar no vermelho'],
    [
        ['Essencial (R$ 119/mes)', 'R$ 40 a 80'],
        ['Profissional (R$ 490/mes)', 'R$ 150 a 300'],
        ['Mix (R$ 250 medio)', 'R$ 100 a 180'],
    ]
)
doc.add_paragraph(
    'CPL real de LinkedIn no mercado financeiro brasileiro: ~R$ 225 a 325. '
    'O Essencial nao comporta midia paga, so organico e ABM direto.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. AGENCIAS PESQUISADAS
# ══════════════════════════════════════════════════════════════
doc.add_heading('5. Pesquisa de agencias qualificadas', level=1)

doc.add_paragraph(
    'Pesquisa realizada em 30/07/2026. Duas das cinco agencias mencionadas no briefing '
    '(BRSA, Seja Mais) nao retornaram resultado verificavel. Abaixo as que foi possivel avaliar.'
)

# 5.1 Sabia
doc.add_heading('Sabia (Agencia Sabia) — Nota 7/10', level=2)
add_tag('[Fato] Fonte: agenciasabia.com.br, BBN International, E3 Network')
add_bullet('Sao Paulo. 16+ anos. Full-service B2B. Membro da BBN (The World\'s B2B Agency) e E3 Network.')
add_bullet('ABM, demand generation, gestao de marca, conteudo, eventos B2B, channel marketing.')
add_bullet('Atlassian: hub de marketing para Brasil/LATAM, evento com 471 inscricoes (6x meta), Golden Bee B2B Awards 2024.')
add_bullet('FECAP Corporate: 105 mil pessoas impactadas, 1.000 acessos, 150 leads qualificados em <30 dias no LinkedIn.')
add_bullet('New Relic: evento "Future Stack" em Sao Paulo.')
add_bullet('Melhores cases B2B verificados entre todas as pesquisadas. Metodo de trabalho transparente.', bold_prefix='Ponto forte: ')
add_bullet('Foco em tecnologia/SaaS, nenhum case em mercado financeiro ou credito.', bold_prefix='Ponto fraco: ')

# 5.2 Intelligenzia
doc.add_heading('Intelligenzia — Nota 6/10', level=2)
add_tag('[Fato] Fonte: intelligenzia.com.br, Semrush, SignalHire')
add_bullet('Fundada em 2012. Sao Paulo. 10-50 funcionarios. Faturamento estimado: US$ 10M-50M (nao verificado).')
add_bullet('Marketing digital B2B, inbound, midia paga, SEO, inside sales, integracao CRM (HubSpot), automacao.')
add_bullet('Cases declarados: Atech, Access, Genesys, Zeiss. NPS 98%. Publica "O Estado do Marketing B2B no Brasil".')
add_bullet('Integracao marketing+vendas+CRM e exatamente o que falta no VIX Radar.', bold_prefix='Ponto forte: ')
add_bullet('Nenhum case no mercado financeiro. Clientes sao industria e tecnologia.', bold_prefix='Ponto fraco: ')

# 5.3 Macfor
doc.add_heading('Macfor — Nota 6/10', level=2)
add_tag('[Fato] Fonte: HubSpot Ecosystem, RemoteRocketShip')
add_bullet('Fundada em 2011. 51-200 funcionarios. HubSpot Solutions Partner.')
add_bullet('Prospeccao outbound multicanal (email, LinkedIn, WhatsApp, telefone, eventos), ABM, agendamento com C-level.')
add_bullet('ABM e outbound multicanal alinhados com a necessidade do VIX Radar.', bold_prefix='Ponto forte: ')
add_bullet('51-200 funcionarios. Piloto de R$ 4.000/mes tende a receber time junior, nao senior.', bold_prefix='Risco: ')

# 5.4 Demais
doc.add_heading('Demais agencias avaliadas', level=2)
add_bullet('Perfil de growth digital, nao de venda consultiva B2B complexa. Case principal e SEO de empresa de educacao.', bold_prefix='Next4 — Nota 5/10. ')
add_bullet('Foco declarado em ABM e outbound. Parceira do Convoy PR Group. Porem: sem site verificavel, sem cases nominais, sem CNPJ. Informacao insuficiente para sustentar avaliacao.', bold_prefix='Conversa.tech — Nota 4/10. ')
add_bullet('Case Finbits (plataforma brasileira de gestao financeira): 285 leads em 3 meses. Porem: agencia do Reino Unido, custo em libra, inviavel para piloto de R$ 15-30 mil.', bold_prefix='Growth Gorilla — Nota 5/10. ')
add_bullet('Nao encontradas em multiplas buscas. Desconsideradas.', bold_prefix='BRSA e Seja Mais. ')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. RANKING E PROBLEMA DO ORCAMENTO
# ══════════════════════════════════════════════════════════════
doc.add_heading('6. Ranking e o problema do orcamento', level=1)

p = doc.add_paragraph()
run = p.add_run('Nenhuma das tres finalistas opera na faixa de R$ 4.000/mes.')
run.bold = True

doc.add_paragraph(
    'Sabia atende Atlassian e New Relic, e membro da BBN. Intelligenzia fatura dezenas de milhoes. '
    'Macfor tem mais de 50 funcionarios. O orcamento do cenario 2 (R$ 21.390 em 90 dias) paga '
    'um mes e meio de fee dessas agencias, nao tres. Se uma agencia qualificada cobrar R$ 8.000/mes, '
    'o piloto vai para R$ 30-35 mil, estourando o teto da premissa de R$ 30 mil.'
)

add_table(
    ['Posicao', 'Agencia', 'Nota', 'Motivo', 'Risco'],
    [
        ['1', 'Sabia', '7/10', 'Cases B2B verificados, ABM premiado, BBN/E3', 'Orcamento incompativel'],
        ['2', 'Intelligenzia', '6/10', 'B2B estabelecida, CRM+inside sales', 'Orcamento incompativel'],
        ['3', 'Macfor', '6/10', 'ABM multicanal, HubSpot Partner', 'Piloto pequeno = time junior'],
    ]
)

doc.add_heading('Fornecedores concretos dentro do orcamento', level=2)

doc.add_heading('Consultores independentes (R$ 3.000 a 6.000/mes)', level=3)
add_bullet('11 anos em prospeccao no LinkedIn. Metodo em 3 pilares. Clientes: 20+ reunioes qualificadas/mes.', bold_prefix='William Tadeu — Prospeccao 10X. ')
add_bullet('Demand Generation. Experiencia em ABM, outbound, automacao. Trabalhou no Efi Bank (2024-2025) com ICP e captacao de contas estrategicas.', bold_prefix='Vanessa Carvalho — Demand Consultoria. ')
add_bullet('16+ anos em marketing e vendas B2B. Consultoria a US$ 60/sessao.', bold_prefix='Gustavo A. (Upwork). ')
add_bullet('LinkedIn Outbound B2B. US$ 18/h. Clientes: Hilton, Marriott, Philips.', bold_prefix='joaoluciano98 (Freelancer). ')

doc.add_heading('Boutiques de outbound (R$ 4.000 a 8.000/mes)', level=3)
add_bullet('Sao Paulo. 2012. Outbound, ABM, CRM, RD Station. 10-50 func. outmarketing.com.br.', bold_prefix='OUTMarketing Brasil. ')
add_bullet('Sao Paulo. 2011. Prospeccao outbound, inbound, marketing data-driven. Tel: +55 11 3765-3003.', bold_prefix='HyTrade. ')
add_bullet('Outbound B2B multicanal (email, LinkedIn, telefone, mala direta). salesboutique.io.', bold_prefix='Sales Boutique. ')

doc.add_heading('SDR as a Service (R$ 2.000 a 5.000/mes)', level=3)
add_bullet('SDR terceirizado a partir de US$ 800/mes (~R$ 4.000). Ativacao em 7 dias. Fully managed.', bold_prefix='Techsho Solutions. ')
add_bullet('SC. 2019. ~24 func. Prospeccao B2B, pre-vendas, SDR. Google for Startups + Endeavor.', bold_prefix='Erah. ')
add_bullet('Campinas/SP. Terceirizacao de prospeccao e consultoria estrategica.', bold_prefix='Sales 3. ')

p = doc.add_paragraph()
run = p.add_run('Recomendacao: ')
run.bold = True
p.add_run('comecar pelos consultores independentes com experiencia em mercado financeiro '
           '(Vanessa Carvalho pelo Efi Bank, William Tadeu pela metodologia). Se nenhum '
           'estiver disponivel, testar Techsho para SDR terceirizado.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. ROTEIRO DE ABORDAGEM
# ══════════════════════════════════════════════════════════════
doc.add_heading('7. Roteiro de abordagem para agencias', level=1)

doc.add_heading('Perguntas de triagem (fazer na primeira conversa)', level=2)
triagem = [
    'Qual foi o ultimo projeto de ABM para publico abaixo de 20 mil pessoas? Mostrar metrica de reuniao, nao de impressao.',
    'Tem case em mercado financeiro, credito, asset ou fintech B2B? Se nao, como adaptam a metodologia?',
    'Quem executaria o piloto no dia a dia? Qual a senioridade? Falar com essa pessoa na segunda reuniao.',
    'Ja temos posicionamento, ICP, mapa competitivo, copy pronta. A proposta parte disso ou refaz tudo? Se refizer, quanto custa?',
    'Qual a menor estrutura de piloto que aceitam? Existe piso de orcamento ou prazo?',
    'Metrica de sucesso do piloto e o que? Se nao houver reuniao com decisor em 90 dias, o que acontece?',
    'Tem clausula de saida sem penalidade? Em quantos dias?',
    'Como cobram? Fee fixo, fee + variavel por reuniao, ou so variavel? Proporcao?',
]
for i, q in enumerate(triagem, 1):
    add_bullet(q, bold_prefix=f'{i}. ')

doc.add_heading('Primeiro contato', level=2)
doc.add_paragraph(
    'O texto exato do primeiro contato (e-mail ou LinkedIn) esta no documento principal '
    '(fase1-plano-comercial-2026-07-30.md, secao 7). Nenhuma mensagem sera enviada a '
    'terceiro sem aprovacao explicita, uma por uma.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 8. BRIEFING
# ══════════════════════════════════════════════════════════════
doc.add_heading('8. Briefing padronizado', level=1)

doc.add_paragraph(
    'Extraido para arquivo proprio: briefing-agencia-vix-radar-2026-07-30.md. '
    'Este arquivo contem o briefing completo (produto, preco, publico, posicionamento, '
    'escopo, restricoes, metricas, modelo de contratacao, anexos) sem informacoes internas '
    'do restante do documento. Pronto para enviar para agencia apos a primeira conversa.'
)

add_bullet('Produto: VIX Radar, 103 emissores, 13 setores, IA + verificacao adversarial, Lei Zero.')
add_bullet('Preco: Essencial R$ 119/mes, Profissional R$ 490/mes. Sem fidelidade.')
add_bullet('Publico: 8-15 mil profissionais no LinkedIn Brasil. Gestores, analistas, tesourarias, family offices.')
add_bullet('Posicionamento: inteligencia acionavel por fracao do custo de terminal.')
add_bullet('Ja existe: posicionamento, ICP, mapa competitivo, copy, anuncios, campanha, briefing visual, landing, apresentacao.')
add_bullet('A executar: ABM, outbound, CRM, funil, midia segmentada, nutricao, metrica.')
add_bullet('NAO faz parte: diagnostico, workshop, redefinicao de ICP, mapeamento competitivo.')
add_bullet('Metrica que vale: reuniao com decisor. Nao vale: impressao, clique, lead generico.')
add_bullet('Modelo: piloto de 90 dias, sem contrato longo, clausula de saida.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. GRADE DE COMPARACAO
# ══════════════════════════════════════════════════════════════
doc.add_heading('9. Grade de comparacao de propostas', level=1)

doc.add_paragraph('Preencher conforme as propostas chegarem. Modelo em branco:')

add_table(
    ['Criterio', 'Sabia', 'Intelligenzia', 'Macfor'],
    [
        ['Prazo minimo', '', '', ''],
        ['Clausula de saida (dias)', '', '', ''],
        ['Aceita piloto 90 dias?', '', '', ''],
        ['ABM com contas nomeadas', '', '', ''],
        ['Outbound (canais)', '', '', ''],
        ['CRM e funil', '', '', ''],
        ['Midia paga LinkedIn', '', '', ''],
        ['Geracao de reuniao com decisor', '', '', ''],
        ['Nutricao e follow-up', '', '', ''],
        ['Exige diagnostico?', '', '', ''],
        ['Fee mensal', '', '', ''],
        ['Fee de setup', '', '', ''],
        ['Variavel por reuniao', '', '', ''],
        ['Variavel por contrato', '', '', ''],
        ['Midia (investimento)', '', '', ''],
        ['Total 90 dias', '', '', ''],
        ['Senioridade do time', '', '', ''],
        ['Case B2B financeiro', '', '', ''],
        ['Case ABM com metrica', '', '', ''],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 10. ORCAMENTO
# ══════════════════════════════════════════════════════════════
doc.add_heading('10. Orcamento estimado total', level=1)

doc.add_heading('Cenario 1 — Piloto com agencia (rota escolhida)', level=2)
add_bullet('ABM, outbound e midia operados por agencia ou consultor. Yan atende as reunioes geradas.')
add_table(
    ['Item', 'Mensal', '90 dias', 'Premissa'],
    [
        ['Fee de agencia/consultor', 'R$ 4.000', 'R$ 12.000', 'Fee mensal, sem setup'],
        ['Midia LinkedIn', 'R$ 2.500', 'R$ 7.500', '~R$ 83/dia, publico estreito'],
        ['Sales Navigator', 'R$ 130', 'R$ 390', '1 licenca'],
        ['CRM (HubSpot starter)', 'R$ 200', 'R$ 600', 'Se nao usar gratuito'],
        ['Ferramentas (Apollo/Clay)', 'R$ 300', 'R$ 900', 'Enriquecimento de leads'],
        ['TOTAL', 'R$ 7.130', 'R$ 21.390', ''],
    ]
)
doc.add_paragraph('Custo por reuniao qualificada: ~R$ 445 (48 em 90 dias). CAC: R$ 2.377.')
p = doc.add_paragraph()
run = p.add_run('O fee de R$ 4.000/mes e o teto para este orcamento. ')
run.bold = True
p.add_run('Se o fornecedor cobrar mais, a conta nao fecha sem aumentar o investimento total.')

doc.add_heading('Cenario 2 — Agressivo (agencia full-service)', level=2)
add_table(
    ['Item', 'Mensal', '90 dias'],
    [
        ['Fee agencia full-service', 'R$ 8.000', 'R$ 24.000'],
        ['Midia LinkedIn', 'R$ 5.000', 'R$ 15.000'],
        ['Sales Navigator (2 licencas)', 'R$ 260', 'R$ 780'],
        ['CRM + ferramentas', 'R$ 800', 'R$ 2.400'],
        ['SDR terceirizado (meio periodo)', 'R$ 3.000', 'R$ 9.000'],
        ['TOTAL', 'R$ 17.060', 'R$ 51.180'],
    ]
)
doc.add_paragraph('So ativar quando houver cobranca funcionando, clientes pagantes e receita que comporte o investimento.')

doc.add_heading('Cenario 3 — Venda solo (referencia)', level=2)
add_table(
    ['Item', 'Mensal', '90 dias'],
    [
        ['Agencia', 'R$ 0', 'R$ 0'],
        ['LinkedIn organico (perfil Yan)', 'R$ 0', 'R$ 0'],
        ['CRM (HubSpot free)', 'R$ 0', 'R$ 0'],
        ['Sales Navigator (opcional)', 'R$ 130', 'R$ 390'],
        ['TOTAL', 'R$ 130', 'R$ 390'],
    ]
)
doc.add_paragraph('Zero agencia. Yan opera o funil sobre rede pessoal. CAC ~R$ 78.')

p = doc.add_paragraph()
run = p.add_run('Nao ativar enquanto nao houver cobranca funcionando e pelo menos 3 clientes pagantes.')
run.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 11. RISCOS
# ══════════════════════════════════════════════════════════════
doc.add_heading('11. Riscos', level=1)

add_bullet('Se confirmado, material comercial que cita "17 usuarios" sem separar pagante de demo precisa ser corrigido.', bold_prefix='Clientes pagantes = 0. ')
add_bullet('Perder o unico case ancora com nome de instituto conhecido. O produto volta a nao ter case publico.', bold_prefix='Mirabaud e contraparte. ')
add_bullet('103 emissores processados com modelo nao-Anthropic. Alegação de verificacao adversarial para esse dia e falsa. Reprocessar.', bold_prefix='Noturno 30/07 contaminado. ')
add_bullet('Chave atual tem prefixo sk-ant- e 108 caracteres, mas nao foi possivel verificar se e nova ou a original rotacionada. Se Task Scheduler nao herdar, matinal de 31/07 falha.', bold_prefix='Chave Anthropic no ambiente. ')
add_bullet('ABM sem Sales Navigator e prospeccao cega. Custo baixo (R$ 130/mes), mas precisa ser ativado.', bold_prefix='Sales Navigator nao contratado. ')
add_bullet('Nenhuma das agencias pesquisadas apresentou case verificavel em mercado financeiro. Qualquer contratacao carrega o risco de a agencia aprender o setor as custas do piloto.', bold_prefix='Zero cases financeiros nas agencias. ')
add_bullet('Orcamento de R$ 4.000/mes nao compra agencia full-service estabelecida. As tres finalistas provavelmente nao aceitam essa faixa.', bold_prefix='Descolamento orcamento vs agencias. ')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 12. DEPENDENCIAS
# ══════════════════════════════════════════════════════════════
doc.add_heading('12. O que depende de voce para destravar a Fase 2', level=1)

deps = [
    'Confirmar ou corrigir as cinco premissas: clientes pagantes, Mirabaud, orcamento, capacidade de reunioes, prazo.',
    'Executar ou delegar os bloqueios P0-1 a P0-9 na ordem listada na secao 1.',
    'Contratar Sales Navigator se a decisao for por ABM.',
    'Aprovar ou rejeitar o texto do primeiro contato com agencias (secao 7 do documento principal).',
    'Dizer se sigo para a Fase 2 (envio dos contatos) ou se ajusta algo na Fase 1.',
]
for i, d in enumerate(deps, 1):
    add_bullet(d, bold_prefix=f'{i}. ')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 13. TESTES DE SENSIBILIDADE
# ══════════════════════════════════════════════════════════════
doc.add_heading('13. Testes de sensibilidade', level=1)

doc.add_heading('O que muda se o ticket cair pela metade', level=2)
doc.add_paragraph(
    'Com R$ 60 (Essencial) e R$ 245 (Profissional), o CPL maximo toleravel cai para R$ 20-40 '
    'no Essencial. Midia paga fica inviavel. A operacao inteira migra para organico e ABM direto. '
    'O cenario com agencia some completamente.'
)

doc.add_heading('O que muda se o orcamento cair pela metade', level=2)
doc.add_paragraph(
    'Com R$ 7.500 a 15.000, so cabe venda solo com ferramentas basicas. '
    'O piloto com agencia ja era apertado em R$ 21.390, na metade nao paga nem um mes de fee.'
)

doc.add_heading('O que muda se clientes pagantes forem zero', level=2)
doc.add_paragraph(
    'A recomendacao de vender sozinho vira absoluta. Contratar agencia com zero pagantes '
    'e sem checkout e queimar dinheiro. Obriga revisao imediata do material comercial.'
)

doc.add_heading('O que nao foi possivel verificar', level=2)
add_bullet('Cases financeiros de qualquer agencia — risco de aprender o setor as custas do piloto.')
add_bullet('Aceite de piloto de 90 dias por agencias — pode ser que recusem e force buscar alternativas menores.')
add_bullet('Preco real de agencia para o escopo — orcamento pode estar subdimensionado.')
add_bullet('Existencia de BRSA, Seja Mais e Conversa.tech — universo de agencias menor que o briefing presumia.')

# ══════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════
output_dir = r'E:\Diretorio\Claude\FREQUENTE\Monitoramento de Credito\marketing\comercial'
output_path = os.path.join(output_dir, 'VIX_Radar_Plano_Comercial_Fase1_2026-07-30.docx')
doc.save(output_path)
print(f'Documento salvo: {output_path}')
